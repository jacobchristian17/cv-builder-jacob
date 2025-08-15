"""DOCX resume parser implementation."""

import logging

logger = logging.getLogger(__name__)


class DOCXParser:
    """Parser for DOCX/DOC resume files."""
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as a string
        """
        try:
            import docx2txt
            return docx2txt.process(file_path)
        except ImportError:
            logger.warning("docx2txt not installed. Trying python-docx...")
            try:
                from docx import Document
                return self._extract_with_python_docx(file_path)
            except ImportError:
                raise ImportError(
                    "No DOCX parsing library found. Install docx2txt or python-docx: "
                    "pip install docx2txt python-docx"
                )
    
    def _extract_with_python_docx(self, file_path: str) -> str:
        """Extract text using python-docx."""
        from docx import Document
        
        doc = Document(file_path)
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return '\n'.join(text)